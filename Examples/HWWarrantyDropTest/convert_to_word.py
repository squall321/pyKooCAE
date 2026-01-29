#!/usr/bin/env python3
"""
Markdown to Word 변환 스크립트
KooChainRun_Complete_Guide.md를 .docx 형식으로 변환합니다.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def parse_markdown_to_word(md_path: str, docx_path: str):
    """
    마크다운 파일을 읽어서 Word 문서로 변환
    """
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    # 코드 스타일 추가
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
    except:
        code_style = doc.styles['Normal']

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_block = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # 코드 블록 처리
        if line.startswith('```'):
            if in_code_block:
                # 코드 블록 종료
                if code_block:
                    code_text = '\n'.join(code_block)
                    p = doc.add_paragraph(code_text, style='Code')
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_block = []
                in_code_block = False
            else:
                # 코드 블록 시작
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block.append(line)
            i += 1
            continue

        # 제목 처리
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)

        # 수평선 처리
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)

        # 리스트 처리
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # 인라인 코드 처리
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # 볼드 처리
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')

        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            text = re.sub(r'`([^`]+)`', r'\1', text)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')

        # 표 처리 (간단한 마크다운 표)
        elif '|' in line and line.strip().startswith('|'):
            # 표 헤더 수집
            table_lines = [line]
            i += 1
            # 구분선 건너뛰기
            if i < len(lines) and '---' in lines[i]:
                i += 1
            # 표 내용 수집
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # 표 생성
            if len(table_lines) >= 2:
                rows = []
                for tline in table_lines:
                    cells = [c.strip() for c in tline.split('|')[1:-1]]
                    rows.append(cells)

                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
            continue

        # 일반 텍스트
        elif line.strip():
            text = line.strip()
            # 인라인 코드 처리
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # 볼드 처리
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # 빈 줄
        else:
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()

        i += 1

    # 저장
    doc.save(docx_path)
    print(f"✅ Word 문서 생성 완료: {docx_path}")
    print(f"   파일 크기: {Path(docx_path).stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    md_file = "KooChainRun_Complete_Guide.md"
    docx_file = "KooChainRun_Complete_Guide.docx"

    parse_markdown_to_word(md_file, docx_file)
